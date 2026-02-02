"""Intelligent document processing with structure awareness."""

import re
import asyncio
import hashlib
import pickle
from concurrent.futures import ThreadPoolExecutor
from typing import List, Optional, Dict, Any, Union
from pathlib import Path
import logging
from dataclasses import dataclass

import PyPDF2
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown
from markdown.extensions import toc, tables, codehilite

from ..models.document import Document, DocumentChunk, DocumentMetadata, DocumentType, ChunkType
from ..config import config

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processes documents with intelligent chunking and structure extraction."""
    
    def __init__(self, max_workers: int = 4, enable_cache: bool = True):
        self.chunk_size = config.retrieval.chunk_size
        self.chunk_overlap = config.retrieval.chunk_overlap
        self.max_workers = max_workers
        self.executor = ThreadPoolExecutor(max_workers=max_workers)
        self.enable_cache = enable_cache and config.system.enable_caching
        self.cache_dir = config.system.cache_dir / "document_processing"
        
        if self.enable_cache:
            self.cache_dir.mkdir(exist_ok=True)
    
    def _get_cache_key(self, file_path: Path) -> str:
        """Generate cache key for a file."""
        stat = file_path.stat()
        content = f"{file_path}_{stat.st_size}_{stat.st_mtime}_{self.chunk_size}_{self.chunk_overlap}"
        return hashlib.md5(content.encode()).hexdigest()
    
    def _get_cached_document(self, cache_key: str) -> Optional[Document]:
        """Retrieve document from cache."""
        if not self.enable_cache:
            return None
        
        cache_file = self.cache_dir / f"{cache_key}.pkl"
        if cache_file.exists():
            try:
                with open(cache_file, 'rb') as f:
                    return pickle.load(f)
            except Exception as e:
                logger.warning(f"Failed to load cache {cache_key}: {e}")
        return None
    
    def _cache_document(self, cache_key: str, document: Document) -> None:
        """Cache processed document."""
        if not self.enable_cache:
            return
        
        cache_file = self.cache_dir / f"{cache_key}.pkl"
        try:
            with open(cache_file, 'wb') as f:
                pickle.dump(document, f)
        except Exception as e:
            logger.warning(f"Failed to cache document {cache_key}: {e}")
        
    async def process_files_batch(self, file_paths: List[Path], 
                                 progress_callback: Optional[callable] = None) -> List[Document]:
        """Process multiple files concurrently."""
        semaphore = asyncio.Semaphore(self.max_workers)
        
        async def process_with_semaphore(file_path: Path, index: int) -> Document:
            async with semaphore:
                try:
                    document = await self.process_file(file_path)
                    if progress_callback:
                        await progress_callback(index + 1, len(file_paths), file_path.name)
                    return document
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    error_doc = Document(metadata=DocumentMetadata(file_path=str(file_path)))
                    error_doc.processing_errors.append(str(e))
                    return error_doc
        
        tasks = [process_with_semaphore(path, i) for i, path in enumerate(file_paths)]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out exceptions and return valid documents
        documents = []
        for result in results:
            if isinstance(result, Document):
                documents.append(result)
            elif isinstance(result, Exception):
                logger.error(f"Batch processing error: {result}")
        
        return documents
        
    async def process_file(self, file_path: Path) -> Document:
        """Process a file and return a Document object."""
        # Check cache first
        cache_key = self._get_cache_key(file_path)
        cached_doc = self._get_cached_document(cache_key)
        if cached_doc:
            logger.info(f"Retrieved from cache: {file_path.name}")
            return cached_doc
        
        # Validation
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.stat().st_size == 0:
            raise ValueError(f"File is empty: {file_path}")
        
        if file_path.stat().st_size > 100 * 1024 * 1024:  # 100MB limit
            raise ValueError(f"File too large (>100MB): {file_path}")
        
        try:
            # Determine document type
            doc_type = self._get_document_type(file_path)
            
            # Extract content based on type
            if doc_type == DocumentType.PDF:
                content, metadata = await self._process_pdf(file_path)
            elif doc_type == DocumentType.DOCX:
                content, metadata = await self._process_docx(file_path)
            elif doc_type == DocumentType.TEXT:
                content, metadata = await self._process_text(file_path)
            elif doc_type == DocumentType.HTML:
                content, metadata = await self._process_html(file_path)
            elif doc_type == DocumentType.MARKDOWN:
                content, metadata = await self._process_markdown(file_path)
            else:
                raise ValueError(f"Unsupported document type: {doc_type}")
            
            # Validate extracted content
            if not content or len(content.strip()) < 10:
                raise ValueError(f"No meaningful content extracted from {file_path}")
            
            # Create document
            document = Document(
                content=content,
                metadata=metadata
            )
            
            # Intelligent chunking
            chunks = await self._create_intelligent_chunks(content, metadata)
            if not chunks:
                raise ValueError(f"No chunks created for {file_path}")
                
            for chunk in chunks:
                document.add_chunk(chunk)
            
            document.is_processed = True
            
            # Cache the processed document
            self._cache_document(cache_key, document)
            
            logger.info(f"Processed document: {file_path.name} ({len(chunks)} chunks)")
            
            return document
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {str(e)}")
            document = Document(metadata=DocumentMetadata(file_path=str(file_path)))
            document.processing_errors.append(str(e))
            return document
    
    def _get_document_type(self, file_path: Path) -> DocumentType:
        """Determine document type from file extension."""
        suffix = file_path.suffix.lower()
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.txt': DocumentType.TEXT,
            '.md': DocumentType.MARKDOWN,
            '.html': DocumentType.HTML,
        }
        return type_mapping.get(suffix, DocumentType.TEXT)
    
    async def _process_pdf(self, file_path: Path) -> tuple[str, DocumentMetadata]:
        """Process PDF with structure extraction."""
        content = ""
        metadata = DocumentMetadata(
            file_path=str(file_path),
            document_type=DocumentType.PDF,
            file_size=file_path.stat().st_size
        )
        
        try:
            # Use PyMuPDF for better text extraction and metadata
            doc = fitz.open(file_path)
            metadata.page_count = len(doc)
            
            # Extract metadata
            pdf_metadata = doc.metadata
            if pdf_metadata.get('title'):
                metadata.title = pdf_metadata['title']
            if pdf_metadata.get('author'):
                metadata.authors = [pdf_metadata['author']]
            
            # Extract text with structure preservation
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                
                # Clean and structure the text
                page_text = self._clean_text(page_text)
                content += f"\n\n--- Page {page_num + 1} ---\n{page_text}"
            
            doc.close()
            
        except Exception as e:
            logger.warning(f"PyMuPDF failed for {file_path}, falling back to PyPDF2: {e}")
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata.page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    page_text = self._clean_text(page_text)
                    content += f"\n\n--- Page {page_num + 1} ---\n{page_text}"
        
        # Extract title and abstract if not found in metadata
        if not metadata.title:
            metadata.title = self._extract_title_from_content(content)
        
        metadata.abstract = self._extract_abstract_from_content(content)
        
        return content, metadata
    
    async def _process_docx(self, file_path: Path) -> tuple[str, DocumentMetadata]:
        """Process DOCX document."""
        doc = DocxDocument(file_path)
        
        metadata = DocumentMetadata(
            file_path=str(file_path),
            document_type=DocumentType.DOCX,
            file_size=file_path.stat().st_size
        )
        
        # Extract content
        content = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content += paragraph.text + "\n\n"
        
        # Extract title from first heading or paragraph
        if doc.paragraphs:
            metadata.title = doc.paragraphs[0].text[:100]
        
        return self._clean_text(content), metadata
    
    async def _process_text(self, file_path: Path) -> tuple[str, DocumentMetadata]:
        """Process plain text file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        metadata = DocumentMetadata(
            file_path=str(file_path),
            document_type=DocumentType.TEXT,
            file_size=file_path.stat().st_size,
            title=file_path.stem
        )
        
        return self._clean_text(content), metadata
    
    async def _process_html(self, file_path: Path) -> tuple[str, DocumentMetadata]:
        """Process HTML document."""
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        metadata = DocumentMetadata(
            file_path=str(file_path),
            document_type=DocumentType.HTML,
            file_size=file_path.stat().st_size
        )
        
        # Extract metadata from HTML
        title_tag = soup.find('title')
        if title_tag:
            metadata.title = title_tag.get_text().strip()
        
        # Look for meta tags
        meta_author = soup.find('meta', attrs={'name': 'author'})
        if meta_author:
            metadata.authors = [meta_author.get('content', '')]
        
        meta_description = soup.find('meta', attrs={'name': 'description'})
        if meta_description:
            metadata.abstract = meta_description.get('content', '')
        
        # Extract text content, preserving structure
        content = ""
        
        # Process headings and paragraphs in order
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'article', 'section']):
            if element.name.startswith('h'):
                content += f"\n\n{'#' * int(element.name[1])} {element.get_text().strip()}\n"
            elif element.get_text().strip():
                content += f"\n{element.get_text().strip()}\n"
        
        return self._clean_text(content), metadata
    
    async def _process_markdown(self, file_path: Path) -> tuple[str, DocumentMetadata]:
        """Process Markdown document."""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
        
        metadata = DocumentMetadata(
            file_path=str(file_path),
            document_type=DocumentType.MARKDOWN,
            file_size=file_path.stat().st_size
        )
        
        # Extract front matter if present
        if md_content.startswith('---'):
            try:
                import yaml
                parts = md_content.split('---', 2)
                if len(parts) >= 3:
                    front_matter = yaml.safe_load(parts[1])
                    md_content = parts[2]
                    
                    # Extract metadata from front matter
                    if isinstance(front_matter, dict):
                        metadata.title = front_matter.get('title')
                        if 'author' in front_matter:
                            authors = front_matter['author']
                            if isinstance(authors, str):
                                metadata.authors = [authors]
                            elif isinstance(authors, list):
                                metadata.authors = authors
                        metadata.abstract = front_matter.get('description') or front_matter.get('abstract')
                        if 'keywords' in front_matter:
                            metadata.keywords = front_matter['keywords']
            except ImportError:
                logger.warning("PyYAML not available, skipping front matter parsing")
            except Exception as e:
                logger.warning(f"Error parsing front matter: {e}")
        
        # Extract title from first heading if not in front matter
        if not metadata.title:
            lines = md_content.split('\n')
            for line in lines[:10]:
                if line.startswith('#'):
                    metadata.title = line.lstrip('#').strip()
                    break
        
        # Use the raw markdown content (preserves structure better than converting to HTML)
        return self._clean_text(md_content), metadata
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove page headers/footers patterns
        text = re.sub(r'\n\d+\n', '\n', text)  # Page numbers
        text = re.sub(r'\n[A-Z\s]{10,}\n', '\n', text)  # Headers in caps
        
        return text.strip()
    
    def _extract_title_from_content(self, content: str) -> Optional[str]:
        """Extract title from document content."""
        lines = content.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if len(line) > 10 and len(line) < 200 and not line.startswith('---'):
                # Likely a title
                return line
        return None
    
    def _extract_abstract_from_content(self, content: str) -> Optional[str]:
        """Extract abstract from document content."""
        # Look for abstract section
        abstract_pattern = r'(?i)abstract[:\s]*\n(.*?)(?=\n\n|\n[A-Z]|\nIntroduction|\n1\.)'
        match = re.search(abstract_pattern, content, re.DOTALL)
        if match:
            return match.group(1).strip()
        return None
    
    async def _create_intelligent_chunks(self, content: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Create intelligent chunks based on document structure."""
        chunks = []
        
        # Split by sections first
        sections = self._split_into_sections(content)
        
        for section_title, section_content in sections:
            # Detect section type for better chunking
            section_type = self._detect_section_type(section_title, section_content)
            
            # Further split large sections
            section_chunks = self._split_section_into_chunks(
                section_content, 
                section_title,
                metadata,
                section_type
            )
            chunks.extend(section_chunks)
        
        # Post-process chunks for quality
        chunks = self._post_process_chunks(chunks)
        
        return chunks
    
    def _detect_section_type(self, title: str, content: str) -> ChunkType:
        """Detect the type of section based on title and content."""
        title_lower = title.lower()
        
        if any(word in title_lower for word in ['abstract', 'summary', 'overview']):
            return ChunkType.SECTION
        elif any(word in title_lower for word in ['table', 'figure', 'chart']):
            return ChunkType.TABLE
        elif title_lower.startswith(('fig', 'figure', 'chart', 'graph')):
            return ChunkType.FIGURE_CAPTION
        elif any(word in title_lower for word in ['header', 'title']):
            return ChunkType.HEADER
        else:
            return ChunkType.PARAGRAPH
    
    def _post_process_chunks(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Post-process chunks to improve quality."""
        processed_chunks = []
        
        for chunk in chunks:
            # Skip very short chunks (unless they're important)
            if len(chunk.content.strip()) < 50 and not chunk.is_title:
                continue
            
            # Merge very short adjacent chunks
            if (processed_chunks and 
                len(chunk.content) < 100 and 
                len(processed_chunks[-1].content) < self.chunk_size - 100 and
                chunk.section_title == processed_chunks[-1].section_title):
                
                # Merge with previous chunk
                processed_chunks[-1].content += "\n\n" + chunk.content
                processed_chunks[-1].end_char = chunk.end_char
                processed_chunks[-1].word_count += chunk.word_count
                processed_chunks[-1].sentence_count += chunk.sentence_count
            else:
                processed_chunks.append(chunk)
        
        return processed_chunks
    
    def _split_into_sections(self, content: str) -> List[tuple[str, str]]:
        """Split content into logical sections."""
        sections = []
        
        # Look for section headers (numbered or capitalized)
        section_pattern = r'\n(?=(?:\d+\.?\s+[A-Z]|[A-Z][A-Z\s]{5,})\n)'
        parts = re.split(section_pattern, content)
        
        current_title = "Introduction"
        for part in parts:
            if part.strip():
                lines = part.split('\n', 2)
                if len(lines) > 1 and self._is_section_header(lines[0]):
                    current_title = lines[0].strip()
                    section_content = '\n'.join(lines[1:]) if len(lines) > 1 else ""
                else:
                    section_content = part
                
                if section_content.strip():
                    sections.append((current_title, section_content.strip()))
        
        return sections if sections else [("Main Content", content)]
    
    def _is_section_header(self, line: str) -> bool:
        """Check if a line is likely a section header."""
        line = line.strip()
        if not line:
            return False
        
        # Numbered sections
        if re.match(r'^\d+\.?\s+[A-Z]', line):
            return True
        
        # All caps headers
        if len(line) > 5 and line.isupper() and ' ' in line:
            return True
        
        # Title case headers
        if len(line) < 100 and line.istitle():
            return True
        
        return False
    
    def _split_section_into_chunks(self, content: str, section_title: str, 
                                  metadata: DocumentMetadata, 
                                  section_type: ChunkType = ChunkType.PARAGRAPH) -> List[DocumentChunk]:
        """Split a section into appropriately sized chunks."""
        chunks = []
        
        # Split by paragraphs first
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        current_chunk = ""
        current_start = 0
        
        for paragraph in paragraphs:
            # Check if adding this paragraph would exceed chunk size
            if len(current_chunk) + len(paragraph) > self.chunk_size and current_chunk:
                # Create chunk
                chunk = DocumentChunk(
                    content=current_chunk.strip(),
                    section_title=section_title,
                    chunk_type=section_type,
                    start_char=current_start,
                    end_char=current_start + len(current_chunk),
                    is_title=section_title.lower() in ['title', 'abstract'],
                    is_abstract=section_title.lower() == 'abstract',
                    contains_citations=self._contains_citations(current_chunk)
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else current_chunk
                current_chunk = overlap_text + "\n\n" + paragraph
                current_start += len(current_chunk) - len(overlap_text) - len(paragraph) - 2
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
        
        # Add final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                content=current_chunk.strip(),
                section_title=section_title,
                chunk_type=section_type,
                start_char=current_start,
                end_char=current_start + len(current_chunk),
                is_title=section_title.lower() in ['title', 'abstract'],
                is_abstract=section_title.lower() == 'abstract',
                contains_citations=self._contains_citations(current_chunk)
            )
            chunks.append(chunk)
        
        return chunks
    
    def _contains_citations(self, text: str) -> bool:
        """Check if text contains citations."""
        citation_patterns = [
            r'\[\d+\]',  # [1], [2], etc.
            r'\([A-Za-z]+,?\s*\d{4}\)',  # (Author, 2023)
            r'et al\.',  # et al.
            r'doi:',  # DOI references
        ]
        
        for pattern in citation_patterns:
            if re.search(pattern, text):
                return True
        return False